#!/usr/bin/env python3
# -*- coding: utf‑8 -*-
"""
word_marker_tools.py
~~~~~~~~~~~~~~~~~~~~
Замена маркеров вида [таблица.поле] в документах Word (.docx),
устойчивая к случаям, когда маркер разбит на несколько run'ов.

Функциональность
----------------
• create_test_document(path)      — генерирует тестовый шаблон с маркерами
• extract_placeholders(path)      — возвращает список всех маркеров в файле
• replace_placeholders(inp, out, mapping) — подменяет маркеры по словарю
• CLI‑режим: создаёт шаблон → печатает маркеры → делает замену

Требуется библиотека `python‑docx`:
    pip install python-docx
"""

import os
import re
from typing import Dict, Iterable, Set, List, Optional, Tuple
import sqlite3

from docx import Document


# ──────────────────────────── helpers ────────────────────────────────────── #
def _rewrite_paragraphs(
    paragraphs: Iterable,
    pattern: re.Pattern,
    mapping: Dict[str, str],
    document = None,  # Документ для создания новых параграфов
) -> None:
    """
    Склеивает текст всех run'ов параграфа, делает замену и,
    если что‑то изменилось, переписывает параграф одним run'ом.
    
    Обрабатывает специальное форматирование для значений с префиксом:
    - LIST: - создает маркированный список
    """
    for paragraph in paragraphs:
        original = "".join(run.text for run in paragraph.runs)
        
        # Проверка на наличие открывающей и закрывающей скобок в параграфе
        if '[' in original and ']' in original:
            # Ищем маркеры со значением LIST: в первую очередь
            list_marker_found = False
            for marker, value in mapping.items():
                marker_pattern = re.escape(f"[{marker}]")
                if re.search(marker_pattern, original) and value.startswith("LIST:") and document:
                    print(f"Найден LIST маркер: [{marker}], обрабатываем как список...")
                    match = re.search(marker_pattern, original)
                    
                    # Получаем текст до и после маркера
                    pre_text = original[:match.start()]
                    post_text = original[match.end():]
                    
                    # Получаем родительский элемент и индекс текущего параграфа
                    parent = paragraph._p.getparent()
                    idx = parent.index(paragraph._p)
                    
                    # Очищаем текущий параграф и добавляем в него текст перед маркером
                    for run in paragraph.runs[::-1]:
                        paragraph._p.remove(run._r)
                    
                    if pre_text:
                        paragraph.add_run(pre_text)
                        # Нужно создать новый параграф для первого элемента списка
                        next_p = document.add_paragraph()
                        next_p_element = next_p._p
                        # Вставляем новый параграф после текущего
                        parent.insert(idx + 1, next_p_element)
                        idx += 1  # Обновляем индекс для дальнейшей вставки
                    else:
                        # Используем текущий параграф для первого элемента списка
                        next_p = paragraph
                    
                    # Разбиваем список элементов
                    list_items = value[5:].split('|')  # Удаляем "LIST:" и разбиваем по разделителю
                    
                    # Для каждого элемента списка создаем параграф с маркером
                    for i, item in enumerate(list_items):
                        if i == 0:
                            # Первый элемент списка идет в параграф next_p
                            current_p = next_p
                        else:
                            # Создаем новый параграф для следующего элемента списка
                            current_p = document.add_paragraph()
                            current_p_element = current_p._p
                            # Вставляем его после предыдущего
                            parent.insert(idx + i, current_p_element)
                        
                        # Очищаем параграф (на всякий случай)
                        if current_p.text:
                            for run in current_p.runs[::-1]:
                                current_p._p.remove(run._r)
                        
                        # Добавляем маркер и текст элемента списка
                        current_p.add_run("• " + item.strip())
                        
                        # Устанавливаем отступ для элементов списка (в EMU - 1/100 точки)
                        current_p.paragraph_format.left_indent = 720000  # 0.5 дюйма = 720000 EMU
                    
                    # Если есть текст после маркера, добавляем его в новый параграф
                    if post_text:
                        last_p = document.add_paragraph()
                        last_p_element = last_p._p
                        parent.insert(idx + len(list_items), last_p_element)
                        last_p.add_run(post_text)
                    
                    list_marker_found = True
                    break  # Обрабатываем только первый найденный LIST: маркер
            
            if list_marker_found:
                continue  # Переходим к следующему параграфу
            
            # Если LIST: маркеров не найдено, выполняем обычную замену
            was_replaced = False
            for marker, value in mapping.items():
                marker_pattern = re.escape(f"[{marker}]")
                if re.search(marker_pattern, original):
                    print(f"Обычная замена маркера: [{marker}] -> {value}")
                    original = re.sub(marker_pattern, value, original)
                    was_replaced = True
            
            if was_replaced:
                # Удаляем все старые run'ы
                for run in paragraph.runs[::-1]:
                    paragraph._p.remove(run._r)
                # Добавляем один новый с замененным текстом
                paragraph.add_run(original)
        else:
            # Если нет маркеров в параграфе, пропускаем его
            continue


def _collect_markers(paragraph):
    """Собирает маркеры из параграфа, даже если они разбиты на несколько runs"""
    text = ""
    markers = set()
    
    # Обновляем регулярное выражение для поддержки нового формата
    pattern = re.compile(r"\[([^\[\]]+)\]")
    
    for run in paragraph.runs:
        text += run.text
        # Ищем маркеры в текущем тексте
        matches = pattern.finditer(text)
        for match in matches:
            markers.add(match.group(1))
    
    return markers


def process_db_markers(db_path: str, markers: List[str]) -> Dict[str, str]:
    """
    Единая функция для обработки маркеров, названия которых совпадают с названиями таблиц,
    а то, что идет после точки, совпадает с названием колонки в этой таблице.
    
    Функция соединяет таблицы по id-шникам внутри колонки и возвращает словарь
    с маркерами и их значениями из базы данных.
    
    Args:
        db_path (str): Путь к файлу базы данных SQLite.
        markers (List[str]): Список маркеров для обработки в формате "таблица.поле".
    
    Returns:
        Dict[str, str]: Словарь, где ключи - маркеры, а значения - данные из базы.
    """
    if not markers:
        return {}

    if not os.path.exists(db_path):
        print(f"Ошибка: База данных не найдена по пути {db_path}")
        return {}
    
    result = {}
    conn = None
    
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        # Получаем информацию о таблицах и их колонках
        tables_info = {}
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = [row[0] for row in cursor.fetchall()]
        
        for table in tables:
            cursor.execute(f"PRAGMA table_info({table})")
            columns = [row[1] for row in cursor.fetchall()]
            tables_info[table] = columns
        
        # Обрабатываем маркеры
        for marker in markers:
            if '.' not in marker:
                continue
                
            table_name, column_name = marker.split('.', 1)
            
            # Проверяем существование таблицы и колонки (с учетом регистра)
            if table_name not in tables_info:
                # Проверяем без учета регистра
                table_match = next((t for t in tables if t.lower() == table_name.lower()), None)
                if not table_match:
                    print(f"Таблица '{table_name}' не найдена в базе данных")
                    continue
                table_name = table_match
            
            columns = tables_info[table_name]
            if column_name not in columns:
                # Проверяем без учета регистра
                column_match = next((c for c in columns if c.lower() == column_name.lower()), None)
                if not column_match:
                    print(f"Колонка '{column_name}' не найдена в таблице '{table_name}'")
                    continue
                column_name = column_match
            
            # Проверяем, является ли колонка внешним ключом (id_*)
            if column_name.startswith('id_'):
                # Пытаемся определить связанную таблицу и получить значение
                related_table = column_name[3:]  # Отбрасываем 'id_'
                
                # Ищем таблицу во множественном числе (если она указана в единственном)
                potential_related_tables = [
                    t for t in tables 
                    if t.lower().startswith(related_table.lower()) or 
                    (t.lower().endswith('ы') and t.lower()[:-1] == related_table.lower())
                ]
                
                if potential_related_tables:
                    related_table = potential_related_tables[0]
                    
                    # Получаем ID записи из основной таблицы
                    cursor.execute(f"SELECT {column_name} FROM {table_name} LIMIT 1")
                    related_id = cursor.fetchone()
                    
                    if related_id and related_id[0]:
                        # Определяем, какое поле использовать из связанной таблицы
                        # Обычно это поле с названием/именем/номером и т.п.
                        display_columns = ['фио', 'название', 'наименование', 'номер', 'имя', 'name', 'title']
                        display_column = next((col for col in tables_info.get(related_table, []) 
                                              if col.lower() in display_columns), None)
                        
                        if display_column:
                            cursor.execute(f"SELECT {display_column} FROM {related_table} WHERE id = ?", (related_id[0],))
                            value = cursor.fetchone()
                            
                            if value:
                                result[marker] = value[0]
                                continue
            
            # Стандартный случай - прямое получение значения из таблицы
            try:
                cursor.execute(f"SELECT {column_name} FROM {table_name} LIMIT 1")
                value = cursor.fetchone()
                
                if value is not None:
                    result[marker] = str(value[0]) if value[0] is not None else ""
                else:
                    print(f"Данные для маркера '{marker}' не найдены")
            except sqlite3.Error as e:
                print(f"Ошибка при получении данных для маркера '{marker}': {e}")
    
    except sqlite3.Error as e:
        print(f"Ошибка базы данных: {e}")
    finally:
        if conn:
            conn.close()
    
    return result


def process_related_tables_markers(db_path: str, main_table_id: Optional[int] = None, markers: Optional[List[str]] = None) -> Dict[str, str]:
    """
    Расширенная функция для обработки маркеров с учетом связанных таблиц.
    
    Args:
        db_path (str): Путь к файлу базы данных SQLite.
        main_table_id (int, optional): ID записи в основной таблице (если известен).
        markers (List[str], optional): Список маркеров для обработки.
        
    Returns:
        Dict[str, str]: Словарь с маркерами и их значениями из базы данных.
    """
    if not markers:
        return {}
        
    if not os.path.exists(db_path):
        print(f"Ошибка: База данных не найдена по пути {db_path}")
        return {}
    
    result = {}
    conn = None
    
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        # Получаем схему базы данных - таблицы, колонки и внешние ключи
        tables_info = {}
        foreign_keys = {}
        
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = [row[0] for row in cursor.fetchall()]
        
        for table in tables:
            # Получаем информацию о колонках
            cursor.execute(f"PRAGMA table_info({table})")
            columns = [row[1] for row in cursor.fetchall()]
            tables_info[table] = columns
            
            # Получаем информацию о внешних ключах
            cursor.execute(f"PRAGMA foreign_key_list({table})")
            table_fks = {}
            for fk in cursor.fetchall():
                from_col = fk[3]  # колонка в текущей таблице
                to_table = fk[2]  # целевая таблица
                to_col = fk[4]    # колонка в целевой таблице
                table_fks[from_col] = (to_table, to_col)
            
            foreign_keys[table] = table_fks
        
        # Обрабатываем маркеры
        for marker in markers:
            if '.' not in marker:
                continue
                
            table_name, column_name = marker.split('.', 1)
            
            # Нормализуем имена таблиц и колонок (поиск без учета регистра)
            table_match = next((t for t in tables if t.lower() == table_name.lower()), None)
            if not table_match:
                print(f"Таблица '{table_name}' не найдена в базе данных")
                continue
            
            table_name = table_match
            columns = tables_info[table_name]
            
            column_match = next((c for c in columns if c.lower() == column_name.lower()), None)
            if not column_match:
                print(f"Колонка '{column_name}' не найдена в таблице '{table_name}'")
                continue
            
            column_name = column_match
            
            # Строим SQL-запрос в зависимости от типа связей
            if main_table_id and table_name in foreign_keys:
                # Ищем запись по идентификатору, если указан основной идентификатор
                query = f"SELECT {column_name} FROM {table_name} WHERE id = ?"
                params = (main_table_id,)
            else:
                # Просто берем первую запись, если не указан конкретный ID
                query = f"SELECT {column_name} FROM {table_name} LIMIT 1"
                params = ()
            
            try:
                cursor.execute(query, params)
                value = cursor.fetchone()
                
                # Если значение - это ID другой таблицы, пытаемся получить связанное значение
                if value and value[0] is not None:
                    # Проверяем, является ли колонка внешним ключом
                    if column_name in foreign_keys.get(table_name, {}):
                        related_table, related_col = foreign_keys[table_name][column_name]
                        
                        # Ищем отображаемую колонку в связанной таблице
                        display_columns = ['фио', 'название', 'наименование', 'номер', 'имя', 'name', 'title']
                        display_column = next(
                            (col for col in tables_info.get(related_table, []) 
                             if col.lower() in display_columns), 
                            related_col  # По умолчанию используем колонку связи
                        )
                        
                        cursor.execute(f"SELECT {display_column} FROM {related_table} WHERE {related_col} = ?", (value[0],))
                        related_value = cursor.fetchone()
                        
                        if related_value:
                            result[marker] = str(related_value[0]) if related_value[0] is not None else ""
                        else:
                            result[marker] = str(value[0])
                    else:
                        result[marker] = str(value[0])
                else:
                    result[marker] = ""
            except sqlite3.Error as e:
                print(f"Ошибка при получении данных для маркера '{marker}': {e}")
    
    except sqlite3.Error as e:
        print(f"Ошибка базы данных: {e}")
    finally:
        if conn:
            conn.close()
    
    return result


# ──────────────────────────── core API ───────────────────────────────────── #
def replace_placeholders(
    input_path: str,
    output_path: str,
    mapping: Dict[str, str],
) -> None:
    """
    Заменяет все маркеры, указанные в *mapping*, во всём документе *input_path*
    и сохраняет результат в *output_path*.

    mapping = { 
        "договоры.номер": "2024‑000001",
        "список_работ(договоры.номер)": "работа 1, работа 2",
        "сумма(договоры.номер)": "15000 руб."
    }
    
    Для маркированного списка используйте специальный префикс "LIST:" в значении.
    Например: {"список_работ(договоры.номер)": "LIST:работа 1|работа 2|работа 3"}
    Каждый элемент списка должен быть разделен символом '|'
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Входной файл не найден: {input_path}")

    if os.path.abspath(input_path) == os.path.abspath(output_path):
        raise ValueError("Входной и выходной пути не могут совпадать")

    out_dir = os.path.dirname(output_path) or "."
    if not os.access(out_dir, os.W_OK):
        raise PermissionError(f"Нет прав на запись в директорию: {out_dir}")

    # общий регекс вида \[(key1|key2|…)\]
    keys_re = "|".join(re.escape(k) for k in mapping)
    pattern = re.compile(r"\[(" + keys_re + r")\]")
    
    # Выводим все ключи для отладки
    print("Ключи для замены:")
    for k in mapping.keys():
        print(f"  - {k}")
    
    # Все маркеры, которые есть в документе
    doc_markers = extract_placeholders(input_path)
    print("Маркеры в документе:")
    for marker in doc_markers:
        print(f"  - {marker}")
        if marker in mapping:
            print(f"    Значение: {mapping[marker]}")
        else:
            print(f"    !!! НЕТ ЗНАЧЕНИЯ В СЛОВАРЕ !!!")

    doc = Document(input_path)

    # тело
    _rewrite_paragraphs(doc.paragraphs, pattern, mapping, doc)

    # таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _rewrite_paragraphs(cell.paragraphs, pattern, mapping, doc)

    # колонтитулы
    for section in doc.sections:
        _rewrite_paragraphs(section.header.paragraphs, pattern, mapping, doc)
        _rewrite_paragraphs(section.footer.paragraphs, pattern, mapping, doc)

    doc.save(output_path)
    print(f"✓ Файл сохранён: {output_path}")


def extract_placeholders(docx_path: str) -> list:
    """
    Возвращает отсортированный список всех уникальных маркеров в документе.
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(docx_path)

    doc = Document(docx_path)
    markers: Set[str] = set()

    # тело
    for paragraph in doc.paragraphs:
        markers.update(_collect_markers(paragraph))

    # таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    markers.update(_collect_markers(paragraph))

    # колонтитулы
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            markers.update(_collect_markers(paragraph))
        for paragraph in section.footer.paragraphs:
            markers.update(_collect_markers(paragraph))

    return sorted(markers)


def create_test_document(output_path: str) -> None:
    """Генерирует тестовый шаблон с разными маркерами"""
    doc = Document()
    
    doc.add_heading("Шаблон Word с маркерами", 0)
    
    paragraph = doc.add_paragraph(
        "Это тестовый документ с набором маркеров для обработки. "
        "Каждый маркер должен быть в квадратных скобках. Некоторые маркеры "
        "могут быть разделены на части из-за форматирования."
    )
    
    paragraph = doc.add_paragraph("Простые маркеры: ")
    paragraph.add_run("[договоры.номер]").bold = True
    paragraph.add_run(" и [договоры.дата] и номер вагона [")
    paragraph.add_run("вагоны.номер").italic = True
    paragraph.add_run("] и подразделение [вагоны.подразделение].")
    
    doc.add_paragraph("Список работ по договору [список_работ(договоры.номер)].")
    doc.add_paragraph("Общая стоимость: [сумма(договоры.номер)].")
    
    # Создадим таблицу
    rows = 3
    cols = 2
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    # Заполним таблицу
    table.cell(0, 0).text = "Параметр"
    table.cell(0, 1).text = "Значение"
    
    table.cell(1, 0).text = "Договор"
    table.cell(1, 1).text = "[договоры.номер]"
    
    table.cell(2, 0).text = "Дата"
    table.cell(2, 1).text = "[договоры.дата]"
    
    doc.save(output_path)
    print(f"✓ Тестовый файл сохранён: {output_path}")


# ───────────────────────────── CLI demo ──────────────────────────────────── #
if __name__ == "__main__":
    here = os.path.dirname(os.path.abspath(__file__))

    # 1. Всегда создаём тестовый шаблон
    template = os.path.join(here, "word.docx")
    create_test_document(template)

    # 2. Ищем маркеры
    vars_found = extract_placeholders(template)
    print("Найдены маркеры:", vars_found)

    # 3. Подмена маркеров
    mapping_demo = {
        "договоры.номер":       "2024‑000001",
        "договоры.дата":        "23.12.2024",
        "вагоны.номер":         "12345",
        "вагоны.подразделение": "ПМС‑5, Свердловская ж/д",
        "список_работ(договоры.номер)": "LIST:Работа 1|Ремонт оси|Покраска",
        "сумма(договоры.номер)": "150000 руб."
    }
    result = os.path.join(here, "result.docx")
    replace_placeholders(template, result, mapping_demo)
